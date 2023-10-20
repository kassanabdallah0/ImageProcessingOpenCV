import os

import matplotlib.pyplot as plt
from docx.shared import RGBColor, Pt
from docx import Document
from docx.shared import Inches
from docx2pdf import convert
import datetime
from fonction_gra import *
import cv2
from scipy.stats import norm
import statistics as sta

def makeworddoc(name, cautp, sym, imorg, quad=False, grav=False, imgName='', taille=(0, 0), cg = None, cq = None, surface_total=0, surface_impatcs = 0, thresh_quad=0, thresh_grav=0, img=None):
    today = datetime.date.today()
    data = []
    im = cv2.imread(imorg)
    if not imgName == '':
        nm = imgName.split('/', -1)
        nmm = nm[-1]
        n = nmm.split('.', -1)
        imgName = 'Plaque_' + n[0]
    cv2.imwrite('img/imageoriginale.jpg', im)
    if os.path.exists('comments/txt/table.bin'):
        try:
            with open('comments/txt/table.bin', 'rb') as file:
                longueur = pickle.load(file)
                i = 0
                while i < longueur:
                    x = pickle.load(file)
                    data.insert(i, x)
                    i += 1
        except (IOError, pickle.UnpicklingError):
            print('Erreur de lecture.')
    if os.path.exists('comments/txt/commenttable.txt'):
        myFile = open("comments/txt/commenttable.txt", "r")
        s = myFile.read()
        myFile.close()
    if os.path.exists('comments/txt/value.bin'):
        try:
            with open('comments/txt/value.bin', 'rb') as file:
                thresh = pickle.load(file)
                filtre = pickle.load(file)
                contours = pickle.load(file)
        except (IOError, pickle.UnpicklingError):
            print('Erreur de lecture.')
    if os.path.exists('txt/horizontaleslider.bin') and os.path.exists('txt/cotationquadrillage.bin'):
        try:
            with open('txt/horizontaleslider.bin', 'rb') as file:
                t_horizontalSlider = pickle.load(file)
                h_horizontalSlider = pickle.load(file)
                v_horizontalSlider = pickle.load(file)
        except (IOError, pickle.UnpicklingError):
                print('Erreur de lecture.')
        try:
            with open('txt/cotationquadrillage.bin', 'rb') as file:
                cautationName = pickle.load(file)
                niveau = pickle.load(file)
                percent = pickle.load(file)
        except (IOError, pickle.UnpicklingError):
                print('Erreur de lecture.')
    if os.path.exists('comments/txt/commentquadrillage.txt'):
        myFile = open("comments/txt/commentquadrillage.txt", "r")
        comentsquadrillage = myFile.read()
        myFile.close()
    if os.path.exists('comments/txt/commentlabel.txt'):
        myFile = open("comments/txt/commentlabel.txt", "r")
        l = myFile.read()
        myFile.close()
    if os.path.exists('comments/txt/value.bin'):
        t = True
        if thresh_grav != 0:
            t = False
        cv2.imwrite('comments/txt/filtreview.jpg', img)
        if os.path.exists("comments/txt/courbe_XY.bin"):
            try:
                with open("comments/txt/courbe_XY.bin", "rb") as file:
                    cx = pickle.load(file)
                    cy = pickle.load(file)
            except (IOError, pickle.UnpicklingError):
                print('Erreur de lecture.')
            if os.path.exists('comments/txt/commentcourbe.txt'):
                myFile = open("comments/txt/commentcourbe.txt", "r")
                graphe = myFile.read()
                myFile.close()
        if os.path.exists("comments/txt/surfacegraphe_XY.bin"):
            try:
                with open("comments/txt/surfacegraphe_XY.bin", "rb") as file:
                    ga = pickle.load(file)
            except (IOError, pickle.UnpicklingError):
                print('Erreur de lecture.')
            if os.path.exists('comments/txt/commentreper.txt'):
                myFile = open("comments/txt/commentreper.txt", "r")
                reper = myFile.read()
                myFile.close()
            area1 = np.sort(ga)
            vaiance = sta.stdev(area1)
            moyen = sta.mean(area1)
            plt.plot(area1, norm.pdf(area1, moyen, vaiance))
            plt.xlabel('Surfaces (mm²)')
            plt.ylabel("nombre d'impacts")
            plt.title("Probabilité de nombre d'impacts / Surfaces")
            plt.savefig('comments/txt/fig.jpg')
            plt.close()
            plt.plot(area1)
            plt.xlabel('impacts')
            plt.ylabel('surface (mm²)')
            plt.title("surface de chaque impact / indice des impacts")
            plt.savefig('comments/txt/figraphe.jpg')
            plt.close()
            plt.title("Répartition des impacts sur la surface totale de l'échantillon (%)")
            tasks = [surface_impatcs, surface_total]
            plt.pie(tasks, labels=['surface des impacts', 'surface totale'], autopct='%1.1f%%')
            plt.axis('equal')
            plt.savefig('comments/txt/pie.jpg')
            plt.close()
    #############
    document = Document()

    header = document.sections[0].header
    hparagraphe = header.add_paragraph()
    logo = hparagraphe.add_run()
    logo.add_picture('comments/txt/logo_HMR.jpg', width=Inches(1))
    hparagraphe.add_run("\n")
    hparagraphe.add_run('\t\t\t\t\t\t\t\t ')
    hparagraphe.add_run('Date:').underline = True
    hparagraphe.add_run(' 'f'{str(today)}')
    header.is_linked_to_previous = False
    document.add_heading('Rapport 'f'{str(imgName)}', 0)

    p = document.add_paragraph('')
    if os.path.exists('comments/txt/value.bin') and grav == True:
        run = p.add_run(f"Cotation:{cautp}_{sym}\n")
        font = run.font
        font.bold = True
        font.color.rgb = RGBColor(255, 0, 0)
        font.size = Pt(16)
        run_1 = p.add_run(f"Le nombre d'impacts : {len(area1)}\n")
        font_1 = run_1.font
        font_1.bold = True
        font_1.color.rgb = RGBColor(0, 0, 165)
        font_1.size = Pt(14)
        p.add_run(f"La méthode utilisée: {cg}\n")
    if os.path.exists('txt/horizontaleslider.bin') and os.path.exists('txt/cotationquadrillage.bin') and quad == True:
        run = p.add_run(f"Cotation:{str(niveau)}\n")
        font = run.font
        font.bold = True
        font.color.rgb = RGBColor(255, 0, 0)
        font.size = Pt(16)
        p.add_run(f"La méthode utilisée: {cq}\n")
    p.add_run(f"taille du rectangle = ({str(taille[0])} mm, {str(taille[1])} mm)")
    '''
    p.add_run('bold').bold = True
    p.add_run(' and some ')
    p.add_run('italic.').italic = True
    '''
    if os.path.exists('comments/txt/value.bin') and grav==True:
        document.add_heading('Gravillonnage', level=0)
        document.add_heading('Filtre', level=1)
        '''
        document.add_paragraph('Intense quote', style='Intense Quote')
        
        document.add_paragraph(
            'first item in unordered list', style='List Bullet'
        )
        document.add_paragraph(
            'first item in ordered list', style='List Number'
        )
        '''
        pp = document.add_paragraph('')
        if thresh_grav != 0:
            pp.add_run('Seuil: ').bold = True
            pp.add_run(str(thresh_grav))
            pp.add_run('      ')
            pp.add_run('Filtre: ').bold = True
            pp.add_run(str(filtre))
            pp.add_run('      ')
            pp.add_run('Contours: ').bold = True
            pp.add_run(str(contours))
        else:
            pp.add_run("\t\t\t\t\t")
            pp.add_run('Filtre automatique').bold = True
        pp.add_run("\n")
        pp.add_run("\n")
        i = pp.add_run()
        i.add_picture('img/img_cut_grav.jpg', width=Inches(2.5))
        i1 = pp.add_run("\t\t")
        i1.add_picture('comments/txt/filtreview.jpg', width=Inches(2.5))
        pp.add_run("\n")
        pp.add_run('Commentaire: ').bold = True
        if os.path.exists('comments/txt/commentlabel.txt'):
            pp.add_run(str(l))
        document.add_page_break()
    '''
    document.add_heading('Analyse', level=1)
    panalyse = document.add_paragraph('')
    print(13)
    table = document.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'numero d impact'
    hdr_cells[1].text = 'X'
    hdr_cells[2].text = 'Y'
    hdr_cells[3].text = 'perimetre'
    hdr_cells[4].text = 'area'
    hdr_cells[5].text = 'pourcentage'
    hdr_cells[6].text = 'cotation'
    for n, x, y, p, a, pour, cot in data:
        row_cells = table.add_row().cells
        row_cells[0].text = n
        row_cells[1].text = x
        row_cells[2].text = y
        row_cells[3].text = p
        row_cells[4].text = a
        row_cells[5].text = pour
        row_cells[6].text = cot
    print(14)
    document.add_paragraph("\n")

    p1 = document.add_paragraph('')
    p1.add_run('Commentaire: ').bold = True
    if os.path.exists('comments/txt/commenttable.txt'):
        p1.add_run(s)
        '''
    if os.path.exists('comments/txt/value.bin') and grav==True:
        document.add_heading('Courbe', level=1)
        pc = document.add_paragraph('')
        courbe = pc.add_run("\t\t\t")
        courbe.add_picture('comments/txt/fig.jpg', width=Inches(3.5))
        pc.add_run("\n")
        graph = pc.add_run("\t\t\t")
        graph.add_picture('comments/txt/figraphe.jpg', width=Inches(3.5))
        pc.add_run("\n")
        pieimg = pc.add_run("\t\t\t")
        pieimg.add_picture('comments/txt/pie.jpg', width=Inches(3.5))
        pc.add_run("\n\n")
        pc.add_run("commentaire:").bold = True
        if os.path.exists('comments/txt/commentcourbe.txt'):
            pc.add_run(str(graphe))
        if os.path.exists('comments/txt/commentreper.txt'):
            pc.add_run("\n")
            pc.add_run(str(reper))
        document.add_page_break()
    if os.path.exists('txt/horizontaleslider.bin') and os.path.exists('txt/cotationquadrillage.bin') and quad==True:
        document.add_heading('QUADRILLAGE', level=0)
        pq = document.add_paragraph('')
        pq.add_run("Norme: ").bold = True
        pq.add_run(str(cautationName))
        pq.add_run('      ')
        pq.add_run("Cotation: ").bold = True
        pq.add_run(str(niveau))
        pq.add_run('      ')
        pq.add_run("Pourcentage: ").bold = True
        pq.add_run(str(percent))
        pq.add_run(' ')
        pq.add_run('%')
        pq.add_run("\n")
        pq.add_run("\n")
        originale = pq.add_run()
        originale.add_picture('img\img_cut_quad.jpg', width=Inches(1.25))
        pq.add_run()
        enleve = pq.add_run("\t\t")
        enleve.add_picture('img\Img_trait.png', width=Inches(1.25))
        pq.add_run()
        contt = pq.add_run("\t\t")
        contt.add_picture('img\QIm_draw.png', width=Inches(1.25))
        pq.add_run("\n")
        pq.add_run("\n")
        if thresh_quad !=0:
            pq.add_run("Seuil: ").bold = True
            pq.add_run(str(t_horizontalSlider))
            pq.add_run('      ')
            pq.add_run("Horizontale: ").bold = True
            pq.add_run(str(h_horizontalSlider))
            pq.add_run('      ')
            pq.add_run("Verticale: ").bold = True
            pq.add_run(str(v_horizontalSlider))
        if thresh_quad == 0:
            pq.add_run("\t\t\t\t\tFiltre automatique").bold = True
        pq.add_run("\n")
        pq.add_run("\n")
        pq.add_run("Commentaire: ").bold = True
        if os.path.exists('comments/txt/commentquadrillage.txt'):
            pq.add_run(str(comentsquadrillage))
        pq.add_run('.')
        pq.add_run("\n")
    porg = document.add_paragraph('')
    i = porg.add_run()
    i.add_picture('img/imageoriginale.jpg', width=Inches(5))
    #document.add_page_break()
    nme = name.split('.',-1)
    document.save(name)

